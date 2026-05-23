# -*- coding: utf-8 -*-
"""build_book.py — the mt-fhnw-agentic *book-export* skill.

Turns DB content sources (the curated `*_EN.md` markdown, with `figspec` blocks)
into professional DOCX books using the bookkit engine. One book = a title plus
an ordered list of chapter sources. Figures are rendered from `figspec` via
`render_figspec.py`; the resulting markdown is converted to bookkit blocks and
laid out with a title page, disclaimer, TOC, per-chapter "Sources & QR codes"
boxes, and a page-referenced index — the same engine that produced the
AI-Norms reference book.

This is implemented as a reusable *function library + driver* (a skill), invoked
like the other framework builders:

  python build_book.py --manifest books.json --src <dir> --tools <dir> --out <dir>

`books.json` = {"books":[{"key","title","subtitle","author","context",
                          "chapters":["Dimension_06_..._EN.md", ...]}]}
Chapter paths are relative to --src. A chapter that is already plain markdown
(no figspec) is passed through unchanged.
"""
import argparse
import json
import os
import re
import subprocess
import sys

import bookkit
from bookkit import new_book, render_blocks, add_page_numbers, _set_font, index_field
from bookkit import HEAD, GREY, INK, HEAD_FONT, BODY_FONT
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------- front matter ----------------
def title_page(doc, book):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(book["title"]); _set_font(r, HEAD_FONT); r.bold = True
    r.font.size = Pt(34); r.font.color.rgb = HEAD
    if book.get("subtitle"):
        p2 = doc.add_paragraph(style="BkSubtitle")
        r2 = p2.add_run(book["subtitle"]); _set_font(r2, HEAD_FONT)
        r2.font.size = Pt(14); r2.font.color.rgb = GREY
    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bookkit._border(pr, edges=("bottom",), color="1F3864", sz=12)
    for _ in range(6):
        doc.add_paragraph()
    pa = doc.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ra = pa.add_run(book.get("author", "Daniel Casota")); _set_font(ra, HEAD_FONT)
    ra.font.size = Pt(16); ra.font.color.rgb = INK
    pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = pc.add_run(book.get("context", "MAS Cybersecurity, IWI, FHNW — May 2026"))
    _set_font(rc, HEAD_FONT); rc.font.size = Pt(11); rc.font.color.rgb = GREY
    doc.add_page_break()


def disclaimer_page(doc):
    for _ in range(2):
        doc.add_paragraph()
    h = doc.add_paragraph(); r = h.add_run("Edition & Provenance")
    _set_font(r, HEAD_FONT); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = HEAD
    paras = [
        "Generated from the mt-fhnw-agentic content store (the single source of truth) by the book-export skill. Every chapter traces to a curated source that passed the framework's English-core, reference, number-verification and figure-standard gates.",
        "This volume is provided for educational and informational purposes. It is not legal, regulatory or professional advice. International standards (ISO/IEC and others) are copyrighted works obtained from their issuing bodies; this book describes and paraphrases, it does not reproduce normative text.",
        "Figures are rendered deterministically from figspec descriptions; links carry QR codes for print. The companion AI Bill-of-Materials & Decision Audit records the AI and human-in-the-loop decisions behind the content.",
        "Research and web links last verified May 2026.",
    ]
    for t in paras:
        pp = doc.add_paragraph(style="Normal")
        rr = pp.add_run(t); _set_font(rr, BODY_FONT); rr.font.size = Pt(9.5); rr.font.color.rgb = GREY
        pp.paragraph_format.space_after = Pt(8)
    doc.add_page_break()


def toc(doc):
    h = doc.add_paragraph(style="BkH1")
    r = h.add_run("Contents"); _set_font(r, HEAD_FONT); r.bold = True
    r.font.size = Pt(22); r.font.color.rgb = HEAD
    p = doc.add_paragraph(); run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = 'Right-click and choose "Update Field" to build the table of contents.'
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, instr, sep, t, e):
        run._r.append(el)
    doc.add_page_break()


# ---------------- markdown -> bookkit blocks ----------------
_IMG = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
_HEAD = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^[-*]\s+(.*)$")
_NUM = re.compile(r"^(\d+)\.\s+(.*)$")
_TABLE_SEP = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")


def _is_table_row(line):
    return line.lstrip().startswith("|") and line.count("|") >= 2


def _split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def md_to_blocks(text, first_page_break=False):
    """Convert one source's markdown to a bookkit block list."""
    lines = text.replace("\r\n", "\n").split("\n")
    blocks = []
    i = 0
    n = len(lines)
    para = []
    bullets = []
    numbered = []
    first_h_seen = not first_page_break

    def flush_para():
        nonlocal para
        if para:
            blocks.append({"p": " ".join(para).strip()})
            para = []

    def flush_bullets():
        nonlocal bullets
        if bullets:
            blocks.append({"bullets": bullets})
            bullets = []

    def flush_numbered():
        nonlocal numbered
        if numbered:
            blocks.append({"numbered": numbered})
            numbered = []

    def flush_all():
        flush_para(); flush_bullets(); flush_numbered()

    # skip a leading YAML front-matter block
    if lines and lines[0].strip() == "---":
        j = 1
        while j < n and lines[j].strip() != "---":
            j += 1
        i = j + 1

    while i < n:
        line = lines[i]
        s = line.strip()

        # skip HTML comments / completion sentinels / horizontal rules
        if s.startswith("<!--") or re.match(r"^[A-Z0-9\-]+-(COMPLETE|RANK-COMPLETE)$", s):
            flush_all(); i += 1; continue
        if s in ("---", "***", "___"):
            flush_all(); blocks.append({"rule": True}); i += 1; continue

        # fenced code / figspec -> skip whole block (figures already rendered)
        if s.startswith("```"):
            flush_all(); i += 1
            while i < n and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        # image
        m = _IMG.match(s)
        if m:
            flush_all()
            blocks.append({"img": m.group(2), "caption": m.group(1), "width_cm": 15})
            i += 1; continue

        # heading
        m = _HEAD.match(s)
        if m:
            flush_all()
            level = min(len(m.group(1)), 4)
            blk = {"h": level, "text": m.group(2).strip()}
            if not first_h_seen and level <= 2:
                blk["page_break"] = True
                first_h_seen = True
            blocks.append(blk)
            i += 1; continue

        # table (header row, separator, then body rows)
        if _is_table_row(line) and i + 1 < n and _TABLE_SEP.match(lines[i + 1]):
            flush_all()
            header = _split_row(line)
            rows = []
            i += 2
            while i < n and _is_table_row(lines[i]):
                rows.append(_split_row(lines[i]))
                i += 1
            blocks.append({"table": {"header": header, "rows": rows}})
            continue

        # bullets / numbered
        m = _BULLET.match(s)
        if m:
            flush_para(); flush_numbered(); bullets.append(m.group(1)); i += 1; continue
        m = _NUM.match(s)
        if m:
            flush_para(); flush_bullets(); numbered.append(m.group(2)); i += 1; continue

        # blockquote -> callout
        if s.startswith(">"):
            flush_all()
            blocks.append({"callout": s.lstrip("> ").strip()})
            i += 1; continue

        # blank line ends paragraph/lists
        if s == "":
            flush_all(); i += 1; continue

        # accumulate paragraph text
        flush_bullets(); flush_numbered()
        para.append(s)
        i += 1

    flush_all()
    return blocks


# ---------------- build one book ----------------
def render_chapter_md(tools_dir, src_path, build_root, stem):
    """Render figspec -> figures + resolved md; return resolved markdown text."""
    out_resolved = os.path.join(build_root, f"{stem}_resolved.md")
    rf = os.path.join(tools_dir, "render_figspec.py")
    try:
        subprocess.run(
            [sys.executable, rf, os.path.abspath(src_path), out_resolved, stem],
            cwd=build_root, check=True, capture_output=True, text=True,
        )
        with open(out_resolved, encoding="utf-8") as f:
            return f.read()
    except Exception:
        # fall back to the raw source (no figures) so the book still builds
        with open(src_path, encoding="utf-8") as f:
            return f.read()


def build_one(book, src_dir, tools_dir, out_dir):
    key = book["key"]
    build_root = os.path.join(out_dir, "_work", key)
    os.makedirs(build_root, exist_ok=True)
    # render_figspec writes figures to BASE/figures/<subdir>/, where BASE is the
    # checkout root three levels above the script (…/<root>/code/tools). bookkit
    # resolves img paths under MEDIA, so MEDIA must equal that same BASE.
    base = os.path.dirname(os.path.dirname(os.path.abspath(tools_dir)))
    bookkit.MEDIA = base
    bookkit.QR_DIR = os.path.join(out_dir, "qr")

    doc = new_book()
    fig = [0]
    idx_seen = set()
    title_page(doc, book)
    disclaimer_page(doc)
    toc(doc)

    for ch in book["chapters"]:
        src_path = os.path.join(src_dir, ch)
        if not os.path.exists(src_path):
            print(f"    ! missing chapter {ch}")
            continue
        stem = re.sub(r"[^A-Za-z0-9]+", "_", os.path.splitext(os.path.basename(ch))[0])
        md = render_chapter_md(tools_dir, src_path, build_root, stem)
        blocks = md_to_blocks(md, first_page_break=True)
        bookkit.reset_links()
        render_blocks(doc, blocks, fig, idx_seen)
        bookkit.flush_sources(doc)

    index_field(doc)
    add_page_numbers(doc)
    out_path = os.path.join(out_dir, f"{key}.docx")
    doc.save(out_path)
    print(f"  + {key}.docx  ({fig[0]} figures, {len(book['chapters'])} chapters)")
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--manifest", required=True)
    ap.add_argument("--src", required=True)
    ap.add_argument("--tools", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--only")
    args = ap.parse_args()
    os.makedirs(args.out, exist_ok=True)
    with open(args.manifest, encoding="utf-8") as f:
        manifest = json.load(f)
    built = []
    for book in manifest["books"]:
        if args.only and book["key"] != args.only:
            continue
        try:
            built.append(build_one(book, args.src, args.tools, args.out))
        except Exception as e:
            print(f"  ! {book['key']} FAILED: {e}")
    print(f"Built {len(built)} book(s) into {args.out}")


if __name__ == "__main__":
    main()
