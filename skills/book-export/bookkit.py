# -*- coding: utf-8 -*-
"""
bookkit.py - Reusable book builder for the AI Norms & Regulations book.

Typography decisions (readability-first, A4 technical non-fiction):
  * Body text:   Georgia 11pt, line spacing 1.32 (multiple), 8pt space-after,
                 justified with hyphenation off -> comfortable long-form reading.
  * Headings:    Calibri (sans, contrast against serif body), dark slate-blue.
                 H1 chapter opener large + page break before.
  * Captions:    Georgia italic 9pt, grey, centred under figures/tables.
  * Lists:       Georgia 11pt, tighter spacing (4pt after).
  * Callouts:    light-grey shaded box, Calibri, for notes/key-points.
  * Page:        A4, 2.2cm side margins, 2.5cm top/bottom; page numbers in footer.

Authoring format: a list of block dicts (see render_blocks for the grammar).
Inline markup inside text: **bold**, *italic*, `code`, [label](url).
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

MEDIA = r"C:\Users\dcaso\book_build\media"

# ---- palette ----
INK        = RGBColor(0x1a, 0x1a, 0x1a)   # near-black body
HEAD       = RGBColor(0x1F, 0x38, 0x64)   # dark slate blue headings
HEAD2      = RGBColor(0x2E, 0x4A, 0x7A)
ACCENT     = RGBColor(0x0B, 0x5C, 0x9E)   # link / accent blue
GREY       = RGBColor(0x66, 0x66, 0x66)
RULE       = "C9D2E0"
CALLOUTBG  = "EEF2F8"
CODEBG     = "F2F2F2"

BODY_FONT = "Georgia"
HEAD_FONT = "Calibri"
MONO_FONT = "Consolas"


def _set_font(run, name):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), name)


def _shade(el, fill):
    pr = el.get_or_add_pPr() if el.tag == qn('w:p') else el
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pr.append(shd)


def _border(p, edges=('top','bottom'), color=RULE, sz=6, space=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for e in edges:
        el = OxmlElement('w:' + e)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), str(space)); el.set(qn('w:color'), color)
        pbdr.append(el)
    pPr.append(pbdr)


# ---------- styles ----------
def build_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    _set_font_default(doc, BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.32
    pf.space_after = Pt(8)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def mk(name, base='Normal'):
        try:
            return styles[name]
        except KeyError:
            from docx.enum.style import WD_STYLE_TYPE
            return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    # Headings
    specs = {
        'BkH1': (22, True, HEAD, HEAD_FONT, 0, 14),
        'BkH2': (16, True, HEAD, HEAD_FONT, 18, 6),
        'BkH3': (13, True, HEAD2, HEAD_FONT, 14, 4),
        'BkH4': (11.5, True, HEAD2, HEAD_FONT, 10, 2),
    }
    for nm, (sz, bold, col, fnt, before, after) in specs.items():
        st = mk(nm)
        st.base_style = styles['Normal']
        st.font.name = fnt; st.font.size = Pt(sz); st.font.bold = bold; st.font.color.rgb = col
        pf = st.paragraph_format
        pf.space_before = Pt(before); pf.space_after = Pt(after)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.05
        pf.keep_with_next = True
        # map to Word outline level so the TOC field and navigation pane pick them up
        lvl = int(nm[-1]) - 1
        pPr = st.element.get_or_add_pPr()
        ol = pPr.find(qn('w:outlineLvl'))
        if ol is None:
            ol = OxmlElement('w:outlineLvl'); pPr.append(ol)
        ol.set(qn('w:val'), str(lvl))

    cap = mk('BkCaption'); cap.base_style = styles['Normal']
    cap.font.name = BODY_FONT; cap.font.size = Pt(9); cap.font.italic = True; cap.font.color.rgb = GREY
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3); cap.paragraph_format.space_after = Pt(12)

    bl = mk('BkBullet'); bl.base_style = styles['Normal']
    bl.paragraph_format.space_after = Pt(4); bl.paragraph_format.left_indent = Cm(0.8)
    bl.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    callout = mk('BkCallout'); callout.base_style = styles['Normal']
    callout.font.name = HEAD_FONT; callout.font.size = Pt(10.5)
    callout.paragraph_format.left_indent = Cm(0.3); callout.paragraph_format.right_indent = Cm(0.3)
    callout.paragraph_format.space_before = Pt(6); callout.paragraph_format.space_after = Pt(6)
    callout.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = mk('BkSubtitle'); sub.base_style = styles['Normal']
    sub.font.name = HEAD_FONT; sub.font.size = Pt(13); sub.font.color.rgb = GREY
    sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER; sub.paragraph_format.space_after = Pt(4)


def _set_font_default(doc, name):
    # set the document default east-asian/ascii font
    el = doc.styles.element
    rpr = el.find(qn('w:docDefaults') + '/' + qn('w:rPrDefault') + '/' + qn('w:rPr'))


# ---------- page setup ----------
def page_setup(doc):
    for s in doc.sections:
        s.page_height = Cm(29.7); s.page_width = Cm(21.0)  # A4
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.3); s.right_margin = Cm(2.3)


def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
        run._r.addnext(fld1)
        _set_font(run, BODY_FONT); run.font.size = Pt(9); run.font.color.rgb = GREY


# ---------- inline markup + link registry (Sources & QR codes) ----------
_INLINE = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))')
QR_DIR = os.path.join(os.path.dirname(MEDIA), 'qr')

# Per-chapter registry of (label, url). Flushed into a 'Sources & QR codes'
# box at the end of each chapter so links are scannable in print.
_LINKS = []


def reset_links():
    _LINKS.clear()


def _register_link(label, url):
    # de-duplicate within a chapter: reuse the number if same url already seen
    for i, (lbl, u) in enumerate(_LINKS, 1):
        if u == url:
            return i
    _LINKS.append((label, url))
    return len(_LINKS)


def _superscript(p, text):
    r = p.add_run(text)
    rpr = r._element.get_or_add_rPr()
    va = OxmlElement('w:vertAlign'); va.set(qn('w:val'), 'superscript'); rpr.append(va)
    _set_font(r, BODY_FONT); r.font.size = Pt(8); r.font.color.rgb = ACCENT
    return r


def add_inline(p, text, bold=True):
    """Parse **bold**, *italic*, `code`, [label](url).

    Links are recorded in the per-chapter registry and rendered as the label
    plus a small superscript reference number; the full URL + a QR code appear
    in the chapter's 'Sources & QR codes' box. When ``bold`` is False, **x**
    renders as plain text (used for running prose so bold stays sparse)."""
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = p.add_run(tok[2:-2]); _set_font(r, BODY_FONT); r.bold = bool(bold)
        elif tok.startswith('`') and tok.endswith('`'):
            r = p.add_run(tok[1:-1]); _set_font(r, MONO_FONT); r.font.size = Pt(10)
            _shade_run(r, CODEBG)
        elif tok.startswith('*') and tok.endswith('*') and len(tok) > 2:
            r = p.add_run(tok[1:-1]); _set_font(r, BODY_FONT); r.italic = True
        elif tok.startswith('['):
            m = re.match(r'\[(.+?)\]\((.+?)\)', tok)
            if m:
                label, url = m.group(1), m.group(2)
                r = p.add_run(label); _set_font(r, BODY_FONT)
                _superscript(p, str(_register_link(label, url)))
            else:
                r = p.add_run(tok); _set_font(r, BODY_FONT)
        else:
            r = p.add_run(tok); _set_font(r, BODY_FONT)


def _shade_run(run, fill):
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill); rpr.append(shd)


def add_hyperlink(paragraph, url, text, size_pt=None, mono=False):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    fnt = MONO_FONT if mono else BODY_FONT
    for a in ('w:ascii','w:hAnsi','w:cs'): rfonts.set(qn(a), fnt)
    rPr.append(rfonts)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0B5C9E'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    if size_pt:
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size_pt * 2))); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    new_run.append(t); hyperlink.append(new_run); paragraph._p.append(hyperlink)
    return hyperlink


def _qr_png(url):
    import hashlib, qrcode
    os.makedirs(QR_DIR, exist_ok=True)
    h = hashlib.md5(url.encode('utf-8')).hexdigest()[:16]
    path = os.path.join(QR_DIR, h + '.png')
    if not os.path.exists(path):
        qr = qrcode.QRCode(error_correction=qrcode.constants.ERROR_CORRECT_M, box_size=10, border=2)
        qr.add_data(url); qr.make(fit=True)
        qr.make_image(fill_color="black", back_color="white").save(path)
    return path


def _qr_row(table, left_lines, url):
    """Add one [text | QR] row to a 2-col sources table.
    left_lines: list of (text, style) where style in {'label','url'}."""
    cells = table.add_row().cells
    lc = cells[0]; lc.text = ''
    first = True
    for txt, style in left_lines:
        p = lc.paragraphs[0] if first else lc.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        if style == 'url':
            add_hyperlink(p, url, txt, size_pt=8, mono=True)
        elif style == 'label':
            r = p.add_run(txt); _set_font(r, BODY_FONT); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = HEAD
        else:
            r = p.add_run(txt); _set_font(r, BODY_FONT); r.font.size = Pt(9)
    rc = cells[1]; rc.text = ''
    pc = rc.paragraphs[0]; pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pc.add_run()
    try:
        run.add_picture(_qr_png(url), width=Cm(2.1))
    except Exception:
        rr = pc.add_run('[QR]'); _set_font(rr, BODY_FONT); rr.font.size = Pt(8)
    cells[0].width = Cm(13.6); cells[1].width = Cm(2.9)


def flush_sources(doc):
    """Emit the chapter's 'Sources & QR codes' box and reset the registry."""
    if not _LINKS:
        return
    # heading styled like H3 but NOT an outline level, so it stays out of the TOC
    h = doc.add_paragraph(style='Normal')
    h.paragraph_format.space_before = Pt(12); h.paragraph_format.keep_with_next = True
    rh = h.add_run('Sources & QR codes'); _set_font(rh, HEAD_FONT); rh.bold = True
    rh.font.size = Pt(12); rh.font.color.rgb = HEAD2
    note = doc.add_paragraph(style='Normal')
    rr = note.add_run('Scan a code to open the link, or read the URL below it.')
    _set_font(rr, BODY_FONT); rr.italic = True; rr.font.size = Pt(9.5); rr.font.color.rgb = GREY
    t = doc.add_table(rows=0, cols=2); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, url) in enumerate(_LINKS, 1):
        _qr_row(t, [(f'{i}.  {label}', 'label'), (url, 'url')], url)
    doc.add_paragraph()
    reset_links()


def _render_qrlink(doc, b):
    """A link-list entry rendered inline: context sentence(s) + URL + QR."""
    label = b.get('label', '')
    ctx = b.get('context', '')
    url = b['url']
    if label or ctx:
        p = doc.add_paragraph(style='Normal'); p.paragraph_format.space_after = Pt(2)
        if label:
            r = p.add_run(label + ('. ' if ctx else '')); _set_font(r, BODY_FONT); r.bold = True
        if ctx:
            add_inline(p, ctx, bold=False)
    t = doc.add_table(rows=0, cols=2); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _qr_row(t, [(url, 'url')], url)
    _table_keep_together(t)
    doc.add_paragraph()


# ---------- index (Word XE / INDEX fields) ----------
# Curated index terms. Each is matched (whole-token, case-insensitive) in body text;
# the first occurrence per chapter is marked with a hidden XE field so Word can
# compile a real, page-referenced index.
INDEX_TERMS = [
    "ISO/IEC 42001", "ISO/IEC 27001", "ISO/IEC 27002", "ISO/IEC 27090", "ISO/IEC 27701",
    "ISO/IEC 23894", "ISO/IEC 23053", "ISO/IEC 22989", "ISO/IEC 5338", "ISO/IEC 5259",
    "ISO/IEC 5962", "ISO/IEC TS 4213", "ISO/IEC 24029", "ISO 9001", "ISO 22301",
    "ISO 31000", "ISO 38500", "ISO 10218", "ISO/TS 15066", "ISO 13482", "ISO 13485",
    "IEC 61508", "SAE J3016", "ETSI EN 303 645", "SPDX", "CycloneDX", "SBOM",
    "EU AI Act", "GDPR", "DSGVO", "FADP", "nDSG", "NIS2", "DORA", "Cyber Resilience Act",
    "Digital Markets Act", "Digital Omnibus", "Cybersecurity Act",
    "NIST", "Cybersecurity Framework", "AI Risk Management Framework", "Risk Management Framework",
    "Privacy Framework", "post-quantum", "adversarial machine learning",
    "ENISA", "OECD", "Cloud Security Alliance", "OWASP", "MITRE", "ATT&CK", "SANS",
    "Secure AI Framework", "SAIF", "CISA", "WIPO", "CEN-CENELEC", "NICE Framework",
    "Switzerland", "European Union", "United States", "United Kingdom", "China",
    "Japan", "South Korea", "ASEAN", "Brazil", "Colombia", "India", "Israel", "Russia",
    "intelligent humanoid", "humanoid", "agentic AI", "AI agent", "prompt injection",
    "data poisoning", "model exfiltration", "Basel III", "export control", "quantum computing",
    "robustness", "explainability", "human oversight", "conformity assessment",
    "Software Bill of Materials", "machine learning", "generative AI", "large language model",
    "FINMA", "FDPIC", "BACS", "Microsoft Purview", "AICM", "drug discovery", "medical device",
    "MDR", "FDA", "model risk", "IKT-Minimalstandard", "automated decision",
    "Council of Europe AI Convention", "GR00T", "Nvidia", "Tesla Optimus", "Unitree",
]
# label -> compiled boundary regex
_IDX = []
for _t in INDEX_TERMS:
    _pat = re.compile(r'(?<![A-Za-z0-9])' + re.escape(_t) + r'(?![A-Za-z0-9])', re.IGNORECASE)
    _IDX.append((_t, _pat))


def _fld_run(paragraph, child):
    """One run carrying a single field child."""
    r = OxmlElement('w:r'); r.append(child); paragraph._p.append(r)


def add_xe(paragraph, term):
    """Insert an index-entry (XE) field as three field runs (begin / instrText / end).
    An XE field has no field *result*, so when field codes are displayed as values
    (the default, and what 'Show paragraph marks' leaves untouched) it shows nothing.
    The INDEX field still compiles it into a real, page-referenced index."""
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = f' XE "{term}" '
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    _fld_run(paragraph, b)
    _fld_run(paragraph, it)
    _fld_run(paragraph, e)


def mark_index(paragraph, text, seen):
    if not text or seen is None:
        return
    for label, pat in _IDX:
        if label in seen:
            continue
        if pat.search(text):
            add_xe(paragraph, label)
            seen.add(label)


def index_field(doc):
    h = doc.add_paragraph(style='BkH1')
    add_inline_head(h, "Index", 'BkH1')
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = ' INDEX \\h "A" \\c "2" \\z "1031" '
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "Right-click and choose “Update Field” to build the index."
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    for el in (b, instr, sep, t, e):
        run._r.append(el)


# ---------- block renderer ----------
def render_blocks(doc, blocks, fig_counter, idx_seen=None):
    """
    Grammar of a block dict:
      {"h":1|2|3|4, "text":..., "page_break":bool}
      {"p":"text with inline markup"}
      {"lead":"text"}                      first paragraph, slightly larger
      {"bullets":[...]}  {"numbered":[...]}
      {"img":"image2.png", "caption":"...", "width_cm":float}
      {"table":{"header":[...], "rows":[[...]]}, "caption":"...", "widths":[...]}
      {"callout":"text", "title":"Key point"}
      {"quote":"text", "by":"..."}
      {"rule":True}
      {"space":True}
    """
    for b in blocks:
        if 'h' in b:
            if b.get('page_break'):
                doc.add_page_break()
            if b['h'] == 1 and idx_seen is not None:
                idx_seen.clear()   # restart index tracking each chapter
            style = f"BkH{b['h']}"
            p = doc.add_paragraph(style=style)
            add_inline_head(p, b['text'], style)
            mark_index(p, b['text'], idx_seen)
        elif 'lead' in b:
            p = doc.add_paragraph(style='Normal')
            add_inline(p, b['lead'], bold=False)
            for r in p.runs:
                if r.font.size is None:
                    r.font.size = Pt(11.5)
            mark_index(p, b['lead'], idx_seen)
        elif 'p' in b:
            p = doc.add_paragraph(style='Normal'); add_inline(p, b['p'], bold=False)
            mark_index(p, b['p'], idx_seen)
        elif 'qrlink' in b:
            _render_qrlink(doc, b['qrlink'])
        elif 'bullets' in b:
            for it in b['bullets']:
                p = doc.add_paragraph(style='BkBullet')
                r = p.add_run('•  '); _set_font(r, BODY_FONT); r.font.color.rgb = ACCENT; r.bold = True
                add_inline(p, it)
                mark_index(p, it, idx_seen)
        elif 'numbered' in b:
            for i, it in enumerate(b['numbered'], 1):
                p = doc.add_paragraph(style='BkBullet')
                r = p.add_run(f'{i}.  '); _set_font(r, BODY_FONT); r.bold = True; r.font.color.rgb = ACCENT
                add_inline(p, it)
                mark_index(p, it, idx_seen)
        elif 'img' in b:
            fig_counter[0] += 1
            path = os.path.join(MEDIA, b['img'])
            if os.path.exists(path):
                pic_p = doc.add_paragraph(); pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pic_p.add_run()
                w = Cm(b.get('width_cm', 14))
                try:
                    run.add_picture(path, width=w)
                except Exception as e:
                    add_inline(pic_p, f"[image {b['img']} could not be embedded]")
                cap = doc.add_paragraph(style='BkCaption')
                add_inline(cap, f"Figure {fig_counter[0]}. {b.get('caption','')}")
        elif 'table' in b:
            _render_table(doc, b, fig_counter)
        elif 'callout' in b:
            _render_callout(doc, b)
        elif 'quote' in b:
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.right_indent = Cm(0.8)
            _border(p, edges=('left',), color='0B5C9E', sz=18, space=10)
            r = p.add_run(b['quote']); _set_font(r, BODY_FONT); r.italic = True
            if b.get('by'):
                p2 = doc.add_paragraph(style='Normal'); p2.paragraph_format.left_indent = Cm(0.8)
                r2 = p2.add_run('— ' + b['by']); _set_font(r2, BODY_FONT); r2.font.color.rgb = GREY; r2.font.size = Pt(10)
        elif 'rule' in b:
            p = doc.add_paragraph(); _border(p, edges=('bottom',))
        elif 'space' in b:
            doc.add_paragraph()


def add_inline_head(p, text, style):
    fnt = HEAD_FONT
    sizes = {'BkH1':22,'BkH2':16,'BkH3':13,'BkH4':11.5}
    for tok in _INLINE.split(text):
        if not tok: continue
        if tok.startswith('**') and tok.endswith('**'):
            r = p.add_run(tok[2:-2])
        elif tok.startswith('`') and tok.endswith('`'):
            r = p.add_run(tok[1:-1])
        else:
            r = p.add_run(tok)
        _set_font(r, fnt)


def _render_callout(doc, b):
    if b.get('title'):
        p = doc.add_paragraph(style='BkCallout')
        _shade(p._p, CALLOUTBG)
        _border(p, edges=('left',), color='1F3864', sz=24, space=8)
        r = p.add_run(b['title']); _set_font(r, HEAD_FONT); r.bold = True; r.font.color.rgb = HEAD; r.font.size = Pt(10.5)
    p = doc.add_paragraph(style='BkCallout')
    _shade(p._p, CALLOUTBG)
    _border(p, edges=('left',), color='1F3864', sz=24, space=8)
    add_inline(p, b['callout'])


def _table_keep_together(t, chain=True):
    """Keep table rows from splitting mid-row; for tables that fit on a page,
    chain rows with keep-with-next so the whole table (and its title) move to
    the next page together rather than stranding the title at a page foot."""
    rows = t.rows
    n = len(rows)
    for ri, row in enumerate(rows):
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement('w:cantSplit'); cant.set(qn('w:val'), 'true'); trPr.append(cant)
        if chain and ri < n - 1:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.keep_with_next = True


def _render_table(doc, b, fig_counter):
    spec = b['table']
    header = spec.get('header')
    rows = spec['rows']
    ncols = len(header) if header else max(len(r) for r in rows)
    # Caption rendered ABOVE the table as a title, kept with the table so the
    # two never separate across a page break.
    cap = b.get('caption')
    if cap:
        cp = doc.add_paragraph(style='BkCaption')
        cp.paragraph_format.keep_with_next = True
        cp.paragraph_format.space_after = Pt(3)
        add_inline(cp, cap)
    t = doc.add_table(rows=0, cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    if header:
        row = t.add_row()
        hr = row.cells
        for i, htext in enumerate(header):
            _cell(hr[i], htext, bold=True, fill='1F3864', color=RGBColor(0xFF,0xFF,0xFF))
        # repeat header row on each page the table spans
        trPr = row._tr.get_or_add_trPr()
        th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)
    for r in rows:
        cells = t.add_row().cells
        for i in range(ncols):
            val = r[i] if i < len(r) else ''
            fill = 'F4F6FA' if (len(t.rows) % 2 == 0) else None
            _cell(cells[i], val, fill=fill)
    # chain keep-with-next only for tables that plausibly fit on one page,
    # so very long reference tables can still break (with header repeat).
    _table_keep_together(t, chain=(len(t.rows) <= 14))
    doc.add_paragraph()


def _cell(cell, text, bold=False, fill=None, color=INK):
    cell.text = ''
    p = cell.paragraphs[0]; p.style = None
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    add_inline(p, str(text))
    for r in p.runs:
        r.font.size = Pt(9.5)
        if bold: r.bold = True
        r.font.color.rgb = color
    if fill:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
        tcPr.append(shd)


def new_book():
    doc = Document()
    build_styles(doc)
    page_setup(doc)
    return doc
